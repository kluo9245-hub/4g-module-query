#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""手动操作指导手册 → Word .docx 生成脚本"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# ============ 标题 ============
h1 = doc.add_heading('4G模块选型工具 — 手动操作指导手册', level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('适用于没有代码基础的用户，按步骤操作即可完成修改并同步到在线网页。')
doc.add_paragraph()

# ============ 一、文件说明 ============
doc.add_heading('一、文件说明', level=1)
doc.add_paragraph('工具包含以下核心文件，知道它们的作用就知道该改哪里：')
table1 = doc.add_table(rows=5, cols=3)
table1.style = 'Table Grid'
headers1 = ['文件', '作用', '什么时候需要改它']
for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        para.runs[0].bold = True

data1 = [
    ['data.json', '存储所有国家/地区及其频段数据', '新增/删除国家、修改某国的频段'],
    ['build.py', '定义所有模块型号、频段、归属产品', '新增/删除模块、修改模块产品型号'],
    ['template.html', '网页的外观和展示逻辑', '一般不需要改，除非要改网页样式'],
    ['index.html', '⚠️ 自动生成，不需要手动改', '每次运行 build.py 后自动更新'],
]
for row_idx, row_data in enumerate(data1, 1):
    for col_idx, text in enumerate(row_data):
        table1.rows[row_idx].cells[col_idx].text = text

doc.add_paragraph()
p_warn = doc.add_paragraph()
run = p_warn.add_run('⚠️ 重要：index.html 是自动生成的，永远不要手动改它！')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# ============ 二、准备工作 ============
doc.add_heading('二、准备工作', level=1)
doc.add_paragraph('用「记事本」或「Notepad++」（推荐）打开文件，不要用 Word。')
doc.add_paragraph('• 记事本：右键文件 → 打开方式 → 记事本')
doc.add_paragraph('• Notepad++：去微软应用商店搜索安装，更友好')

# ============ 三、修改国家数据 ============
doc.add_heading('三、修改国家数据（data.json）', level=1)
doc.add_paragraph('用记事本打开 data.json，内容格式如下：')
p_code1 = doc.add_paragraph()
p_code1.add_run('{\n  "countries": [\n    ["中国", "亚太", "FDD B1/B3/B5/B8 TDD B34/B38/B39/B40/B41", "综合评估说明文字"],\n    ["日本", "亚太", "FDD B1/B3/B8/B11/B21 TDD B41", "..."],\n    ...\n  ]\n}').font.name = 'Consolas'
doc.add_paragraph()
doc.add_paragraph('每一个行代表一个国家，四个部分用逗号分隔：')
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
headers2 = ['位置', '说明']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
for para in table2.rows[0].cells[0].paragraphs:
    para.runs[0].bold = True
for para in table2.rows[0].cells[1].paragraphs:
    para.runs[0].bold = True
data2 = [
    ['第1个', '国家/地区名称，如 "中国"'],
    ['第2个', '所属区域，如 "亚太"、"欧洲"、"北美" 等'],
    ['第3个', '4G频段，格式：FDD Bx/By TDD Bz'],
    ['第4个', '综合评估说明文字，可任意填写'],
]
for row_idx, row_data in enumerate(data2, 1):
    table2.rows[row_idx].cells[0].text = row_data[0]
    table2.rows[row_idx].cells[1].text = row_data[1]

doc.add_heading('新增国家', level=2)
doc.add_paragraph('在 "countries" 的方括号 [ ] 里，最后一个国家后面加逗号，然后新增一行：')
p_code2 = doc.add_paragraph()
p_code2.add_run('["泰国", "亚太", "FDD B1/B3/B8 TDD B41", "东南亚常用频段..."]').font.name = 'Consolas'
p_warn2 = doc.add_paragraph()
run2 = p_warn2.add_run('⚠️ 注意：逗号是英文逗号 , 不是中文逗号 ，')
run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
run2.bold = True

doc.add_heading('删除国家', level=2)
doc.add_paragraph('找到对应国家的那一行，整行删除（包括结尾的逗号）。')
doc.add_heading('修改频段', level=2)
doc.add_paragraph('修改第3个字段，格式固定为：FDD Bx/By TDD Bz，频段之间用 / 分隔。')

# ============ 四、修改模块和产品型号 ============
doc.add_heading('四、修改模块和产品型号（build.py）', level=1)
doc.add_paragraph('用记事本打开 build.py，找到 MODULES = { 这一段（大约在文件第28行）。')

doc.add_heading('模块结构说明', level=2)
p_code3 = doc.add_paragraph()
p_code3.add_run(
    "'CN': {\n"
    "    'name': 'MC669-CN / LE270-CN / ML307N-DC-DL',\n"
    "    'short_name': 'CN三合一',\n"
    "    'bands': {'FDD': ['B1','B3','B5','B8'], 'TDD': ['B34','B38','B39','B40','B41']},\n"
    "    'has_gsm': False,\n"
    "    'priority': 100,\n"
    "    'region': '亚太/中东/非洲/大洋洲',\n"
    "    'note': '不支持GSM，适合亚太地区',\n"
    "    'projects': {'MC669-CN': ['WD-219G'], 'LE270-CN': ['WD-300','WD-210']}\n"
    "}"
).font.name = 'Consolas'

doc.add_heading('修改产品型号（projects 字段）', level=2)
doc.add_paragraph('projects 有两种格式：')
doc.add_paragraph('格式A：简单列表（一个模块对应多个产品）', style='List Bullet')
p_code4 = doc.add_paragraph()
p_code4.add_run("'projects': ['WD-300', 'WD-210', 'WD-219K']").font.name = 'Consolas'
doc.add_paragraph('格式B：按子型号分组（CN三合一、EU版用这种）', style='List Bullet')
p_code5 = doc.add_paragraph()
p_code5.add_run(
    "'projects': {\n"
    "    'MC669-CN': ['WD-219G'],\n"
    "    'LE270-CN': ['WD-300', 'WD-210', 'WD-218', 'WD-219K'],\n"
    "    'ML307N-DC-DL': ['WD-110', 'WD-281', 'WD-282', 'MW-100', 'WD-280D']\n"
    "}"
).font.name = 'Consolas'
doc.add_paragraph('• 新增产品型号：在列表里加 \'产品名\',（注意逗号）')
doc.add_paragraph('• 删除产品型号：把对应 \'产品名\' 整行删掉（包括逗号）')

doc.add_heading('新增模块', level=2)
doc.add_paragraph('在 MODULES = { 和最后一个 } 之间，复制一个现有模块的代码块，修改对应内容，注意：')
doc.add_paragraph('• 新模块的开头键名（如 \'NEW\'）不能和现有模块重复', style='List Bullet')
doc.add_paragraph('• 每个模块之间用逗号 , 分隔', style='List Bullet')
doc.add_paragraph('• priority 数字越大排在越前面', style='List Bullet')

doc.add_heading('删除模块', level=2)
doc.add_paragraph('找到要删除的模块代码块（从开头键名到结尾的 }），整块删除，注意前后的逗号要处理干净。')

# ============ 五、生成网页 ============
doc.add_heading('五、生成网页（运行 build.py）', level=1)
doc.add_paragraph('改完 data.json 或 build.py 之后，必须运行 build.py 重新生成 index.html。')

doc.add_heading('方法一：双击运行（推荐）', level=2)
steps1 = [
    '双击 build.py 文件',
    '会弹出一个黑色窗口，显示类似内容：',
    '  ✅ 从 data.json 读取 171 个国家',
    '  📊 自动匹配完成：共 821 条国家-模块关联',
    '  ✨ 构建完成！',
    '看到 ✨ 构建完成！就说明成功',
    '按任意键关闭窗口',
]
for step in steps1:
    doc.add_paragraph(step, style='List Number')
p_note1 = doc.add_paragraph()
run_note1 = p_note1.add_run('如果双击没有反应，用方法二。')
run_note1.italic = True

doc.add_heading('方法二：命令行运行', level=2)
doc.add_paragraph('1. 按 Win + R，输入 cmd，回车')
doc.add_paragraph('2. 输入以下命令（路径根据实际情况调整）：')
p_code6 = doc.add_paragraph()
p_code6.add_run(
    '"C:\\Users\\NME\\.workbuddy\\binaries\\python\\versions\\3.13.12\\python.exe"'
    ' "C:\\Users\\NME\\WorkBuddy\\20260420143552\\4g_module_tool\\build.py"'
).font.name = 'Consolas'
doc.add_paragraph('3. 看到 ✨ 构建完成！ 就说明成功')

# ============ 六、推送到 GitHub ============
doc.add_heading('六、推送到 GitHub（让在线网页更新）', level=1)
doc.add_paragraph('只有完成这步，别人访问你的在线网页才能看到最新内容。')

doc.add_heading('方法一：用 GitHub Desktop（最简单，推荐）', level=2)
steps_desktop = [
    '去 https://desktop.github.com/ 下载安装 GitHub Desktop',
    '安装后登录你的 GitHub 账号',
    '找到 4g-module-query 仓库，点击 Clone',
    '以后每次修改完文件，GitHub Desktop 会自动检测到变化',
    '在左下角填写 Summary（随便写，如"更新国家数据"）',
    '点击 Commit to main 按钮',
    '再点击右上角 Push origin 按钮',
    '等待几分钟，在线网页自动更新',
]
for step in steps_desktop:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('方法二：用命令行（已配置好，直接复制粘贴）', level=2)
doc.add_paragraph('1. 打开 4g_module_tool 文件夹')
doc.add_paragraph('2. 在文件夹地址栏输入 cmd 回车，打开命令行')
doc.add_paragraph('3. 依次输入以下命令：')
p_code7 = doc.add_paragraph()
p_code7.add_run(
    '# 第1步：查看哪些文件被修改了\n'
    'git status\n\n'
    '# 第2步：把所有修改的文件加入提交\n'
    'git add data.json build.py template.html index.html\n\n'
    '# 第3步：提交（引号里写本次修改的说明）\n'
    "git commit -m \"更新数据\"\n\n"
    '# 第4步：推送到GitHub（最关键的一步）\n'
    'git push'
).font.name = 'Consolas'
doc.add_paragraph('4. 看到类似 main -> main 的字样就说明推送成功')
doc.add_paragraph('5. 等待 1-3 分钟，访问网页就能看到更新')

# ============ 七、常见问题 ============
doc.add_heading('七、常见问题', level=1)
faqs = [
    ('运行 build.py 报错怎么办？', '检查 data.json 是否有格式错误，最常见的是少了逗号或多了逗号。把报错截图发给我。'),
    ('git push 需要输入用户名密码怎么办？', '说明你还没有配置 GitHub 凭证，告诉我，我帮你配置。'),
    ('我不小心改错了文件怎么办？', '右键文件 → 属性 → 以前的版本（如果有备份），或者告诉我，我可以帮你恢复。'),
    ('推送后网页没有立即更新？', 'GitHub Pages 需要 1-3 分钟部署，稍等再刷新页面（Ctrl+F5 强制刷新）。'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run_q = p.add_run('Q：' + q)
    run_q.bold = True
    doc.add_paragraph('A：' + a)

# ============ 八、操作检查清单 ============
doc.add_heading('八、操作检查清单', level=1)
doc.add_paragraph('每次修改后，按这个顺序检查：')
checklist = [
    '修改了 data.json 或 build.py',
    '运行了 build.py 并看到 ✨ 构建完成！',
    '双击 index.html 打开本地网页，确认内容正确',
    '执行了 git push 推送到 GitHub',
    '等待 1-3 分钟后在线网页已更新',
]
for item in checklist:
    doc.add_paragraph('□ ' + item, style='List Bullet')

# ============ 保存 ============
output_path = 'C:/Users/NING MEI/WorkBuddy/20260420143552/4g_module_tool/4G模块选型工具-手动操作指导手册.docx'
doc.save(output_path)
print('✅ 已生成：' + output_path)
